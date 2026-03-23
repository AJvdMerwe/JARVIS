"""
tools/writing_tools.py
────────────────────────
Tools for the WritingAssistantAgent.

  OutlineTool        — generate a structured outline from a brief
  DraftSectionTool   — write one section (heading + body) at a time
  AssembleDraftTool  — join all sections into a complete draft
  EditDraftTool      — improve grammar, clarity, style, or length
  ExportMarkdownTool — save the final draft as a Markdown file
  ExportDocxTool     — export the final draft as a formatted .docx file
  ListDraftsTool     — list drafts stored in the current session
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from langchain_core.tools import BaseTool
from pydantic import BaseModel, Field

from config import settings

logger = logging.getLogger(__name__)

# ── In-process draft registry ─────────────────────────────────────────────────
# Maps draft_id → {"title": str, "sections": list[dict], "full_text": str}
_DRAFTS: dict[str, dict] = {}

_WRITING_DIR = settings.uploads_path / "generated" / "writing"
_WRITING_DIR.mkdir(parents=True, exist_ok=True)

_SAFE_RE = re.compile(r"[^\w\-]")


def _safe_name(s: str, ext: str) -> str:
    stem = _SAFE_RE.sub("_", Path(s).stem.strip())[:60] or "document"
    return f"{stem}{ext}"


# =============================================================================
#  Schemas
# =============================================================================

class OutlineInput(BaseModel):
    topic:        str = Field(..., description="The topic or title of the document.")
    brief:        str = Field("", description="Additional context, requirements, or key points to cover.")
    doc_type:     str = Field(
        "article",
        description="Document type: article | essay | report | blog_post | technical_doc | email | letter",
    )
    target_words: int = Field(800, ge=100, le=10000, description="Approximate target word count.")
    draft_id:     str = Field("draft", description="Short ID to store this draft under.")


class DraftSectionInput(BaseModel):
    draft_id:     str = Field("draft", description="Draft ID (from create_outline).")
    section_idx:  int = Field(0, ge=0, description="0-based index of the section to write.")
    extra_notes:  str = Field("", description="Additional instructions for this specific section.")


class AssembleDraftInput(BaseModel):
    draft_id: str = Field("draft", description="Draft ID to assemble into a full document.")


class EditDraftInput(BaseModel):
    draft_id:     str = Field("draft", description="Draft ID to edit.")
    instructions: str = Field(
        ...,
        description=(
            "Editing instructions. Examples: "
            "'Make it more concise', "
            "'Fix grammar and punctuation', "
            "'Add more technical detail to section 2', "
            "'Improve the conclusion', "
            "'Adjust tone to be more formal'."
        ),
    )
    section_idx:  Optional[int] = Field(
        None, description="If set, edit only this section (0-based). Leave blank to edit the whole draft."
    )


class ExportMarkdownInput(BaseModel):
    draft_id: str = Field("draft", description="Draft ID to export.")
    filename: str = Field("document.md", description="Output filename.")


class ExportDocxInput(BaseModel):
    draft_id: str = Field("draft", description="Draft ID to export.")
    filename: str = Field("document.docx", description="Output filename.")
    style:    str = Field(
        "professional",
        description="Document style: professional | academic | casual",
    )


# =============================================================================
#  Tools
# =============================================================================

class OutlineTool(BaseTool):
    """
    Generate a structured document outline from a topic brief.

    Creates a hierarchical outline with section headings and bullet-point
    notes for each section. Stores the outline in the draft registry under
    the given draft_id for subsequent section drafting.
    """
    name: str        = "create_outline"
    description: str = (
        "Generate a structured outline for a document. Call this first before drafting. "
        "Inputs: topic (str), brief (str, optional extra context), "
        "doc_type (article|essay|report|blog_post|technical_doc|email|letter), "
        "target_words (int, default 800), draft_id (str, default 'draft')."
    )
    args_schema: type[BaseModel] = OutlineInput

    def _run(
        self,
        topic:        str,
        brief:        str  = "",
        doc_type:     str  = "article",
        target_words: int  = 800,
        draft_id:     str  = "draft",
        **kw,
    ) -> str:
        from core.llm_manager import get_llm

        prompt = (
            f"Create a detailed outline for a {doc_type} about: {topic}\n"
            + (f"Key requirements: {brief}\n" if brief else "")
            + f"Target length: approximately {target_words} words\n\n"
            "Return a JSON object with this structure:\n"
            '{"title": "...", "sections": [{"heading": "...", "notes": "..."}]}\n'
            "Include 4-8 sections appropriate for the document type. "
            "Notes should be 1-2 sentences describing what to cover. "
            "Return ONLY valid JSON, no markdown fences."
        )

        try:
            import json
            llm    = get_llm()
            result = llm.invoke(prompt)
            raw    = str(result.content).strip()
            raw    = re.sub(r"^```json?\s*|\s*```$", "", raw, flags=re.MULTILINE).strip()
            data   = json.loads(raw)

            _DRAFTS[draft_id] = {
                "title":    data.get("title", topic),
                "topic":    topic,
                "doc_type": doc_type,
                "sections": [
                    {"heading": s["heading"], "notes": s.get("notes", ""), "body": ""}
                    for s in data.get("sections", [])
                ],
                "full_text": "",
                "target_words": target_words,
            }

            lines = [f"Outline for: {data.get('title', topic)}\n"]
            for i, s in enumerate(_DRAFTS[draft_id]["sections"]):
                lines.append(f"  [{i}] {s['heading']}")
                if s["notes"]:
                    lines.append(f"       → {s['notes']}")
            lines.append(f"\nDraft ID: '{draft_id}'  |  {len(_DRAFTS[draft_id]['sections'])} sections")
            lines.append("Use draft_section to write each section, then assemble_draft.")
            return "\n".join(lines)

        except Exception as exc:
            return f"Outline generation failed: {exc}"

    async def _arun(self, **kw) -> str:
        raise NotImplementedError


class DraftSectionTool(BaseTool):
    """
    Write one section of the document at a time.

    Uses the outline notes and previous sections as context so each section
    flows naturally from the last. Stores the written body in the draft registry.
    """
    name: str        = "draft_section"
    description: str = (
        "Write one section of the document. Call create_outline first. "
        "Inputs: draft_id (str), section_idx (int, 0-based), extra_notes (str, optional)."
    )
    args_schema: type[BaseModel] = DraftSectionInput

    def _run(
        self,
        draft_id:    str = "draft",
        section_idx: int = 0,
        extra_notes: str = "",
        **kw,
    ) -> str:
        from core.llm_manager import get_llm

        draft = _DRAFTS.get(draft_id)
        if not draft:
            return f"Draft '{draft_id}' not found. Call create_outline first."
        sections = draft["sections"]
        if section_idx >= len(sections):
            return f"Section index {section_idx} out of range (0–{len(sections)-1})."

        section    = sections[section_idx]
        prev_bodies = "\n\n".join(
            f"## {sections[i]['heading']}\n{sections[i]['body']}"
            for i in range(section_idx)
            if sections[i]["body"]
        )

        words_each = max(80, draft["target_words"] // max(len(sections), 1))
        prompt = (
            f"You are writing a {draft['doc_type']} titled: \"{draft['title']}\"\n\n"
            + (f"Sections written so far:\n{prev_bodies}\n\n" if prev_bodies else "")
            + f"Now write section: **{section['heading']}**\n"
            + (f"Section notes: {section['notes']}\n" if section["notes"] else "")
            + (f"Additional instructions: {extra_notes}\n" if extra_notes else "")
            + f"\nWrite approximately {words_each} words of flowing prose. "
            "Do not include the heading — write only the body text. "
            "Maintain consistent tone with previous sections."
        )

        try:
            llm    = get_llm()
            result = llm.invoke(prompt)
            body   = str(result.content).strip()
            sections[section_idx]["body"] = body
            word_count = len(body.split())
            return (
                f"Section [{section_idx}] '{section['heading']}' written "
                f"({word_count} words).\n\n"
                f"{body[:300]}{'…' if len(body) > 300 else ''}"
            )
        except Exception as exc:
            return f"Section drafting failed: {exc}"

    async def _arun(self, **kw) -> str:
        raise NotImplementedError


class AssembleDraftTool(BaseTool):
    """
    Combine all drafted sections into a single Markdown document.

    Any sections without a written body are skipped with a note.
    The assembled text is stored in the draft registry as `full_text`.
    """
    name: str        = "assemble_draft"
    description: str = (
        "Join all drafted sections into a complete Markdown document. "
        "Call this after draft_section for all sections. "
        "Inputs: draft_id (str, default 'draft')."
    )
    args_schema: type[BaseModel] = AssembleDraftInput

    def _run(self, draft_id: str = "draft", **kw) -> str:
        draft = _DRAFTS.get(draft_id)
        if not draft:
            return f"Draft '{draft_id}' not found. Call create_outline first."

        lines = [f"# {draft['title']}\n"]
        skipped = []
        total_words = 0

        for i, sec in enumerate(draft["sections"]):
            if sec["body"].strip():
                lines.append(f"## {sec['heading']}\n")
                lines.append(sec["body"])
                lines.append("")
                total_words += len(sec["body"].split())
            else:
                skipped.append(f"[{i}] {sec['heading']}")

        full_text = "\n".join(lines)
        draft["full_text"] = full_text

        summary = [
            f"Draft assembled: '{draft['title']}'",
            f"  Total words : {total_words:,}",
            f"  Sections    : {len(draft['sections']) - len(skipped)}/{len(draft['sections'])} written",
        ]
        if skipped:
            summary.append(f"  Skipped     : {', '.join(skipped)}")
        summary.append("\nUse edit_draft to refine, or export_docx / export_markdown to save.")
        return "\n".join(summary)

    async def _arun(self, **kw) -> str:
        raise NotImplementedError


class EditDraftTool(BaseTool):
    """
    Edit the full draft or a single section based on instructions.

    Applies LLM-powered editing: grammar fixes, tone adjustment, length
    changes, clarity improvements, restructuring, or any natural-language
    instruction. Updates the draft registry in place.
    """
    name: str        = "edit_draft"
    description: str = (
        "Edit the full draft or a specific section. "
        "Inputs: draft_id (str), instructions (str, what to change), "
        "section_idx (int, optional — omit to edit the whole draft)."
    )
    args_schema: type[BaseModel] = EditDraftInput

    def _run(
        self,
        draft_id:     str           = "draft",
        instructions: str           = "",
        section_idx:  Optional[int] = None,
        **kw,
    ) -> str:
        from core.llm_manager import get_llm

        draft = _DRAFTS.get(draft_id)
        if not draft:
            return f"Draft '{draft_id}' not found."

        if section_idx is not None:
            # Edit a single section
            sections = draft["sections"]
            if section_idx >= len(sections):
                return f"Section {section_idx} out of range."
            original = sections[section_idx]["body"]
            if not original.strip():
                return f"Section [{section_idx}] has no body yet. Draft it first."
            target_label = f"section [{section_idx}] '{sections[section_idx]['heading']}'"
            text_to_edit = original
        else:
            # Edit the full draft
            if not draft["full_text"].strip():
                return "Draft has no full text. Call assemble_draft first."
            target_label = "full draft"
            text_to_edit = draft["full_text"]

        prompt = (
            f"Edit the following {draft['doc_type']} {target_label}.\n\n"
            f"Instructions: {instructions}\n\n"
            f"Original text:\n{text_to_edit}\n\n"
            "Return ONLY the edited text, preserving Markdown headings where present. "
            "Do not add explanatory commentary."
        )

        try:
            llm    = get_llm()
            result = llm.invoke(prompt)
            edited = str(result.content).strip()

            if section_idx is not None:
                draft["sections"][section_idx]["body"] = edited
            else:
                draft["full_text"] = edited

            original_words = len(text_to_edit.split())
            edited_words   = len(edited.split())
            delta          = edited_words - original_words
            sign           = "+" if delta >= 0 else ""
            return (
                f"Edited {target_label}.\n"
                f"  Before: {original_words:,} words\n"
                f"  After : {edited_words:,} words  ({sign}{delta})\n\n"
                f"{edited[:400]}{'…' if len(edited) > 400 else ''}"
            )
        except Exception as exc:
            return f"Edit failed: {exc}"

    async def _arun(self, **kw) -> str:
        raise NotImplementedError


class ExportMarkdownTool(BaseTool):
    """Save the assembled draft as a Markdown file."""
    name: str        = "export_markdown"
    description: str = (
        "Save the assembled draft as a .md Markdown file. "
        "Inputs: draft_id (str), filename (str, e.g. 'my_article.md')."
    )
    args_schema: type[BaseModel] = ExportMarkdownInput

    def _run(self, draft_id: str = "draft", filename: str = "document.md", **kw) -> str:
        draft = _DRAFTS.get(draft_id)
        if not draft:
            return f"Draft '{draft_id}' not found."
        text = draft.get("full_text") or "\n\n".join(
            f"## {s['heading']}\n{s['body']}" for s in draft["sections"] if s["body"]
        )
        if not text.strip():
            return "Nothing to export — assemble_draft first."
        path = _WRITING_DIR / _safe_name(filename, ".md")
        path.write_text(text, encoding="utf-8")
        words = len(text.split())
        logger.info("Markdown exported: %s", path)
        return f"Saved: {path}\n({words:,} words, {path.stat().st_size:,} bytes)"

    async def _arun(self, **kw) -> str:
        raise NotImplementedError


class ExportDocxTool(BaseTool):
    """
    Export the assembled draft as a formatted Microsoft Word (.docx) file.

    Applies heading styles (H1 for the title, H2 for section headings),
    normal paragraph style for body text, and a professional document theme.
    """
    name: str        = "export_docx"
    description: str = (
        "Export the assembled draft as a formatted .docx Word file. "
        "Inputs: draft_id (str), filename (str, e.g. 'report.docx'), "
        "style (professional|academic|casual)."
    )
    args_schema: type[BaseModel] = ExportDocxInput

    def _run(
        self,
        draft_id: str = "draft",
        filename: str = "document.docx",
        style:    str = "professional",
        **kw,
    ) -> str:
        draft = _DRAFTS.get(draft_id)
        if not draft:
            return f"Draft '{draft_id}' not found."

        full_text = draft.get("full_text", "")
        if not full_text.strip():
            return "Nothing to export — assemble_draft first."

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # ── Page margins ───────────────────────────────────────────────
            for section in doc.sections:
                section.top_margin    = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin   = Inches(1.25)
                section.right_margin  = Inches(1.25)

            # ── Style configuration ────────────────────────────────────────
            normal = doc.styles["Normal"]
            normal.font.name = "Calibri" if style == "professional" else (
                "Times New Roman" if style == "academic" else "Arial"
            )
            normal.font.size = Pt(11)

            # ── Parse and write Markdown content ───────────────────────────
            for line in full_text.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue

                if stripped.startswith("# "):
                    p = doc.add_heading(stripped[2:], level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith(("- ", "* ", "+ ")):
                    # Bullet point
                    p = doc.add_paragraph(stripped[2:], style="List Bullet")
                elif re.match(r"^\d+\. ", stripped):
                    p = doc.add_paragraph(
                        re.sub(r"^\d+\. ", "", stripped), style="List Number"
                    )
                else:
                    # Normal paragraph — handle **bold** and *italic*
                    p = doc.add_paragraph()
                    _write_inline(p, stripped)

            path = _WRITING_DIR / _safe_name(filename, ".docx")
            doc.save(str(path))

            words = len(full_text.split())
            logger.info("DOCX exported: %s", path)
            return f"Saved: {path}\n({words:,} words, {path.stat().st_size:,} bytes)"

        except ImportError:
            return "python-docx is required: pip install python-docx"
        except Exception as exc:
            return f"DOCX export failed: {exc}"

    async def _arun(self, **kw) -> str:
        raise NotImplementedError


def _write_inline(paragraph, text: str) -> None:
    """Write text with **bold** and *italic* Markdown inline markers as Word runs."""
    from docx.shared import Pt
    # Split on bold/italic markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


class ListDraftsTool(BaseTool):
    """List all drafts currently in the writing session."""
    name: str        = "list_drafts"
    description: str = "List all document drafts in the current session. No inputs required."

    def _run(self, **kw) -> str:
        if not _DRAFTS:
            return "No drafts in session. Use create_outline to start a new document."
        lines = ["Active drafts:"]
        for did, d in _DRAFTS.items():
            written = sum(1 for s in d["sections"] if s["body"].strip())
            total   = len(d["sections"])
            assembled = "assembled" if d.get("full_text") else f"{written}/{total} sections"
            lines.append(f"  '{did}'  — {d['title'][:60]}  [{assembled}]")
        return "\n".join(lines)

    async def _arun(self, **kw) -> str:
        raise NotImplementedError


def get_writing_tools() -> list[BaseTool]:
    """Return all writing-assistant tools."""
    return [
        OutlineTool(),
        DraftSectionTool(),
        AssembleDraftTool(),
        EditDraftTool(),
        ExportMarkdownTool(),
        ExportDocxTool(),
        ListDraftsTool(),
    ]
